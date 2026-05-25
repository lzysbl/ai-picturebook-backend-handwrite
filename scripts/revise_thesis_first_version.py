"""Revise the first thesis version for content/format feedback.

The script edits the existing DOCX in place after making a timestamped backup.
It focuses on content boundaries (theory/design/implementation/testing),
chapter titles, abstract metrics, Chinese parentheses, and repeated table
headers for cross-page tables.
"""

from __future__ import annotations

import re
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "报告"
DOC_PATH = REPORT_DIR / "论文第一版_姚家路_20225973.docx"


TITLE_REPLACEMENTS = {
    "3 系统需求分析": "3 绘本讲述系统需求分析",
    "3.1 用户需求": "3.1 用户与使用场景需求",
    "3.2 功能需求": "3.2 功能需求分析",
    "3.3 非功能需求": "3.3 非功能需求分析",
    "4 系统总体设计": "4 多模态绘本讲述系统总体设计",
    "4.1 系统架构设计": "4.1 系统总体架构设计",
    "4.2 数据库设计": "4.2 数据模型与实体关系设计",
    "4.3 实时识别模块设计": "4.3 实时识别与连续讲述流程设计",
    "4.4 响应模式设计": "4.4 多响应模式设计",
    "4.5 语音朗读模块设计": "4.5 语音朗读流程设计",
    "5 系统详细实现": "5 系统功能模块实现",
    "5.1 后端接口实现": "5.1 后端接口与业务流程实现",
    "5.2 多模态模型调用实现": "5.2 多模态识别与讲述生成实现",
    "5.3 实时扫描会话与缓存实现": "5.3 连续识别会话管理实现",
    "5.4 前端实时识别实现": "5.4 前端交互与移动端适配实现",
    "5.5 基础质量约束与启发式评价实现": "5.5 质量约束与日志统计实现",
    "6 系统测试与综合评价": "6 系统测试与定量评价",
    "6.1 测试环境": "6.1 测试环境与数据集",
    "6.2 响应延迟测试": "6.2 识别响应延迟评价",
    "6.3 日志统计结果": "6.3 运行日志指标统计",
    "6.4 人工主观评价": "6.4 讲述质量人工评价",
    "6.5 结果分析": "6.5 测试结论与问题分析",
}


ABSTRACT_CN = [
    (
        "绘本阅读是一种典型的图文结合阅读活动，儿童在阅读过程中不仅需要理解页面中的角色、"
        "场景和动作，还需要将连续页面之间的情节联系起来。随着多模态大模型在图像理解和"
        "自然语言生成任务中的发展，利用模型对绘本页面进行理解并生成适龄讲述文本成为可能。"
        "针对移动端逐页拍摄绘本时存在的图片整理成本高、连续讲述不方便、语音反馈不足等问题，"
        "本文设计并实现了一个基于多模态大模型的绘本讲述应用，支持绘本管理、图片上传、故事生成、"
        "实时识别、连续页保留、故事历史和语音朗读等功能。"
    ),
    (
        "系统以手机或电脑摄像头采集的绘本页面图像以及用户上传的绘本图片为输入，围绕“当前页即时讲述”"
        "和“多页连续讲述”两类场景组织功能流程。普通绘本生成用于整本绘本的离线故事生成，实时识别用于"
        "逐页拍摄时的快速反馈；连续识别通过会话标识保留页级摘要，使后续页面生成时能够参考前文。"
        "语音模块负责将当前页或总故事文本转化为可播放音频，从而形成图像识别、文本讲述和语音反馈结合的"
        "辅助阅读流程。"
    ),
    (
        "本文从功能测试、响应延迟、日志统计和人工复核四个方面进行验证。标准样本选取 demo_book 中"
        "28 本绘本的 55 张页面，流式直接讲述方案识别成功率为 55/55，成功率达到 100%，平均端到端延迟"
        "为 5.672 秒，平均首字返回延迟为 4.976 秒；手机实拍样本选取《小猫钓鱼》12 张连续页面，识别"
        "成功率为 12/12，成功率达到 100%，平均端到端延迟为 8.394 秒，平均首字返回延迟为 7.829 秒。"
        "完整生成模式在抽样测试中的平均延迟为 45.080 秒，更适合作为保存完整故事的高质量模式；"
        "端到端抽样中 Edge TTS 平均合成耗时为 1.714 秒，能够满足短文本朗读需求。"
    ),
    (
        "实验结果表明，系统能够在典型绘本页面和手机实拍页面上稳定完成识别与讲述，其中流式直接讲述"
        "在实时反馈和等待时间之间取得较好平衡，适合作为移动端实时阅读的默认方案；完整生成模式虽然"
        "耗时较长，但适合用于完整故事归档。由此可见，将多模态识别、连续上下文管理和语音合成结合，"
        "可以为绘本辅助阅读提供具有可用性的工程实现路径。"
    ),
]


ABSTRACT_EN = [
    (
        "Picture-book reading is a typical multimodal activity that combines visual scenes and short texts. "
        "Children need to understand characters, actions, emotions, and the narrative relations between "
        "consecutive pages. With the development of multimodal large models, it has become feasible to "
        "interpret picture-book pages and generate age-appropriate storytelling text. To reduce the cost of "
        "manual image preparation and support page-by-page mobile reading, this thesis designs and implements "
        "a multimodal picture-book storytelling application."
    ),
    (
        "The system accepts images captured by a mobile phone or computer camera as well as uploaded page "
        "images. It supports book management, image upload, story generation, real-time page recognition, "
        "continuous page retention, story history, and text-to-speech playback. The normal generation mode is "
        "used for complete story creation, while the real-time recognition mode provides immediate feedback for "
        "the current page. A session-based context mechanism keeps page-level summaries so that later pages can "
        "refer to previous narrative information."
    ),
    (
        "The system is evaluated through functional testing, latency measurement, runtime log statistics, and "
        "manual review. In the standard dataset, 55 pages selected from 28 picture books in demo_book are all "
        "successfully processed by the direct streaming mode, achieving a success rate of 55/55. The average "
        "end-to-end latency is 5.672 seconds, and the average first-token latency is 4.976 seconds. In the "
        "mobile-captured dataset, 12 consecutive pages of Little Cat Fishing are all successfully processed, "
        "with an average latency of 8.394 seconds and an average first-token latency of 7.829 seconds. The full "
        "generation mode has an average latency of 45.080 seconds and is more suitable for archived complete "
        "stories. The sampled Edge TTS synthesis latency is 1.714 seconds."
    ),
    (
        "The results show that the system can complete picture-book recognition and storytelling on both "
        "standard page images and mobile-captured images. Direct streaming storytelling provides a practical "
        "balance between response speed and narrative quality, while full generation is suitable for higher "
        "quality story preservation. The combination of multimodal recognition, continuous context management, "
        "and speech synthesis provides a feasible engineering path for assisted picture-book reading."
    ),
]


THEORY_REPLACEMENTS = {
    "多模态大模型可以把图像、文本等不同输入放到同一任务里处理。以视觉语言模型为例，通常先由视觉编码器提取图像里的对象、场景和空间关系，再通过跨模态对齐模块转成语言模型能使用的表示，最后由语言模型生成回答、说明或故事。本文使用的云端模型大体也遵循这样的思路，只是具体结构由服务商封装完成。": (
        "多模态大模型是指能够同时处理图像、文本、语音等多种模态信息的大规模模型。以视觉语言模型为例，其基本思想是先从图像中提取对象、场景、空间关系和文字等视觉线索，再将这些线索与语言表示对齐，最后生成自然语言回答、说明或故事。与单一模态模型相比，多模态大模型能够在同一任务中综合视觉内容和文本指令，因此更适合处理绘本这类图文共同表达意义的材料。"
    ),
    "本系统没有重新训练本地视觉模型，而是直接调用云端多模态 API。这样做能减少本地显卡和部署压力，也更符合本科毕业设计原型验证的条件。系统把普通绘本生成和实时识别分成两套配置，是为了在生成质量和响应延迟之间留出调整空间。": (
        "从原理上看，多模态大模型的优势在于开放式理解和生成。传统图像分类模型通常只能输出固定类别，图像描述模型多生成客观说明，而多模态大模型可以根据提示词控制回答角度、语言风格和输出长度。绘本页面包含角色、动作、背景、文字和情绪，开放式生成能力可以帮助模型从单纯“看见物体”进一步转向“组织讲述”。"
    ),
    "多模态模型在绘本场景中也会出错。单页识别时，模型可能把背景人物当成主角，也可能在文字模糊时进行不可靠推断。绘本页面同时包含插图、少量文字和版式结构，这和文档智能中的版面理解、OCR、图文联合建模有相似之处，LayoutLMv2、LayoutLMv3、Donut 和 TrOCR 等研究为这类复杂页面理解提供了参考[24-27]。本文通过引导框裁剪、提示词约束、连续页管理和重新识别入口，提高系统可控性。": (
        "绘本页面理解也具有一定复杂性。页面中既有插图，也可能包含标题、对白、页码和版式信息；手机拍摄时还可能出现倾斜、遮挡、反光和背景干扰。这类问题与文档智能中的版面理解、OCR 和图文联合建模有关，LayoutLMv2、LayoutLMv3、Donut 和 TrOCR 等研究为复杂页面理解提供了参考[24-27]。"
    ),
    "在系统里，多模态模型主要干两件事。第一是从页面图片里找出角色、场景、物体、文字和情绪线索；第二是把这些线索组织成适合儿童听读的讲述文本。接口层做成兼容 OpenAI 风格的调用方式，后续更换供应商或比较不同模型时不需要大改业务代码。": (
        "不过，多模态大模型仍然存在幻觉和不稳定问题。模型可能将背景元素误认为故事主体，也可能在图像模糊时生成不存在的情节。因此，在绘本讲述任务中，需要通过输入质量控制、提示约束和结果复核来降低误识别带来的影响。"
    ),
    "连续页管理是绘本讲述必须处理的问题。用户连续拍摄多页时，如果每次识别都覆盖上一次文本，最后只能得到当前页内容。本文把实时扫描结果分为“当前页文本”和“总故事文本”：当前页用于即时查看和朗读，总故事用于累积多页内容并保存到历史记录。": (
        "连续页关系是绘本讲述区别于单图描述的重要特征。用户连续阅读多页时，角色和场景通常具有延续性，后续页面的讲述需要参考前文，否则容易出现角色称呼变化、情节突然断裂或重复讲述等问题。因此，绘本讲述需要在单页理解之外引入上下文保持机制。"
    ),
    "本文把生成文本设计为短篇讲述，而不是识别报告。低龄儿童更适合具体、简短、可听的表达；年龄稍高时，可以适当增加角色心理和情节线索。系统在前端提供目标年龄和讲述风格选择，后端根据这些配置组合提示词。": (
        "从儿童阅读角度看，绘本讲述还需要考虑适龄性。低龄儿童更适合具体、简短、可听的表达；年龄稍高时，可以适当增加角色心理和情节线索。讲述文本不应只是识别报告，而应在事实基础上形成温和、连贯、容易听懂的短篇叙事。"
    ),
    "FastAPI 是基于 Python 的高性能 Web 框架，有类型注解友好、自动生成接口文档和异步处理能力强等特点。本文后端采用 FastAPI 构建 REST 接口和流式接口，使用 Pydantic 定义请求与响应模型，保证接口参数校验和返回结构一致。": (
        "FastAPI 是基于 Python 的现代 Web 框架，具有类型注解友好、自动接口文档和异步处理能力强等特点。对于多模态应用而言，后端通常需要同时处理文件上传、模型调用、数据库访问和流式返回等任务，异步 Web 框架可以提高 I/O 密集型场景下的资源利用率。"
    ),
    "数据库访问使用 SQLAlchemy Async。开发阶段用 SQLite 启动最方便，部署或联调时可以切到 MySQL。Redis 主要用于保存扫描结果缓存和会话上下文，避免同一页面短时间内反复调用模型。": (
        "在 Web 系统中，关系型数据库常用于保存用户、资源、任务和历史记录等结构化数据，缓存系统则常用于保存短期状态、会话上下文和高频访问结果。对于需要连续交互的应用，缓存可以降低重复请求成本，并为多轮上下文管理提供临时存储基础。"
    ),
    "实时识别时，用户最明显感受到的是等待。如果系统必须等整段文本都生成完再显示，页面会空一段时间。本文在实时扫描里加入 SSE：后端把讲述文本按块返回，前端边收边显示，并记录第一块文本到达时间 `first_delta_ms`。": (
        "Server-Sent Events 是一种基于 HTTP 的单向流式通信机制，适合服务器持续向浏览器推送文本块、状态更新或日志消息。与一次性返回完整结果相比，流式输出可以缩短用户看到第一段内容前的等待时间，提升实时交互场景中的感知速度。"
    ),
    "这里把首块返回延迟记为 `first_delta_ms = t_first_chunk_arrive - t_request_sent`。其中 `t_request_sent` 是前端提交识别请求的时间，`t_first_chunk_arrive` 是第一段 SSE 文本块到达的时间。相比完整返回时间，它更接近用户看到第一句讲述前的等待时间。": (
        "在流式系统评价中，除了总响应时间，还常关注首块返回时间。首块返回时间指从用户发起请求到第一段可读结果出现之间的时间，它更接近用户主观感受到的等待时间。对于实时讲述应用，首块返回时间比完整文本返回时间更能反映交互体验。"
    ),
    "语音朗读是绘本辅助阅读中比较直接的功能。本文把 TTS 封装为可替换服务，先后尝试 Bark、Piper 和 Edge TTS。Edge TTS 实际是对 Microsoft Edge 在线神经语音合成服务的程序化调用：系统先清洗和分段待朗读文本，再设置中文语音、语速和音调，最后把返回音频保存为可播放文件。对于本文的短文本和中短篇故事朗读场景，Edge TTS 在中文自然度、部署成本和响应速度之间比较均衡。": (
        "语音合成技术用于将文本转化为自然语音，是绘本辅助阅读中连接文本生成与听觉反馈的重要环节。早期 TTS 系统通常依赖拼接或统计声学模型，近年来神经网络语音合成逐渐成为主流。Tacotron 2、FastSpeech 2、VITS 等方法分别从序列到序列建模、并行生成和端到端建模角度提升了语音自然度与生成效率[30-32]。"
    ),
    "对于长文本，系统先按标点和最大长度切成若干段，再逐段合成音频，最后把结果保存到 `uploads/tts` 并返回播放地址。这样当前页和总故事都可以朗读，也能避开单次合成长度限制。Tacotron 2、FastSpeech 2、VITS 等工作分别从声学建模、并行生成和端到端合成角度推动了 TTS 技术发展，为本文选择在线轻量方案提供了背景[30][31][32]。": (
        "在应用层面，TTS 方案需要同时考虑语音自然度、中文效果、响应延迟、部署成本和调用稳定性。绘本讲述场景中的文本通常较短，但总故事文本可能较长，因此语音合成系统还需要具备分段处理和连续播放能力。"
    ),
}


def contains_cjk(text: str) -> bool:
    return any("\u4e00" <= ch <= "\u9fff" for ch in text)


def normalize_chinese_parentheses(text: str) -> str:
    if not contains_cjk(text):
        return text
    text = re.sub(r"\((\d+)\)", r"（\1）", text)
    text = re.sub(r"\(([\u4e00-\u9fffA-Za-z0-9，、；：。/ ]{1,40})\)", r"（\1）", text)
    return text


def set_paragraph_text(paragraph, text: str) -> None:
    paragraph.text = text


def replace_exact_paragraphs(doc: Document, replacements: dict[str, str]) -> int:
    changed = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in replacements:
            set_paragraph_text(paragraph, replacements[text])
            changed += 1
    return changed


def replace_title_text(doc: Document) -> int:
    changed = 0
    for paragraph in doc.paragraphs:
        original = paragraph.text
        updated = original
        for old, new in TITLE_REPLACEMENTS.items():
            updated = updated.replace(old, new)
        if updated != original:
            paragraph.text = updated
            changed += 1
    return changed


def replace_section_after_heading(doc: Document, heading: str, new_paragraphs: list[str]) -> int:
    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == heading:
            current = index + 1
            marker_index = None
            targets = []
            while current < len(doc.paragraphs):
                text = doc.paragraphs[current].text.strip()
                if text.startswith("关键词") or text == "ABSTRACT" or text.startswith("Keywords:"):
                    marker_index = current
                    break
                targets.append(current)
                current += 1
            if marker_index is None:
                raise RuntimeError(f"section marker not found after {heading}")
            while len(targets) < len(new_paragraphs):
                inserted = doc.add_paragraph("")
                doc.paragraphs[marker_index]._p.addprevious(inserted._p)
                targets.append(marker_index)
                marker_index += 1
            if len(targets) < len(new_paragraphs):
                raise RuntimeError(f"not enough paragraphs after {heading}")
            for target, new_text in zip(targets, new_paragraphs):
                doc.paragraphs[target].text = new_text
            # Clear extra old abstract paragraphs until the next marker.
            current = targets[-1] + 1
            while current < len(doc.paragraphs):
                text = doc.paragraphs[current].text.strip()
                if text.startswith("关键词") or text == "ABSTRACT" or text.startswith("Keywords:"):
                    break
                if text:
                    doc.paragraphs[current].text = ""
                current += 1
            return len(targets)
    raise RuntimeError(f"heading not found: {heading}")


def apply_parentheses_to_doc(doc: Document) -> int:
    changed = 0
    for paragraph in doc.paragraphs:
        updated = normalize_chinese_parentheses(paragraph.text)
        if updated != paragraph.text:
            paragraph.text = updated
            changed += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    updated = normalize_chinese_parentheses(paragraph.text)
                    if updated != paragraph.text:
                        paragraph.text = updated
                        changed += 1
    return changed


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_tables_for_cross_page(doc: Document) -> int:
    changed = 0
    # Skip front-matter metadata tables. Body tables start after the first two.
    for table in doc.tables[2:]:
        if table.rows:
            set_repeat_table_header(table.rows[0])
            changed += 1
    return changed


def main() -> None:
    if not DOC_PATH.exists():
        raise FileNotFoundError(DOC_PATH)
    backup = REPORT_DIR / f"论文第一版_姚家路_20225973_backup_before_feedback_{datetime.now():%Y%m%d_%H%M%S}.docx"
    shutil.copy2(DOC_PATH, backup)

    doc = Document(DOC_PATH)
    abstract_cn_count = replace_section_after_heading(doc, "摘  要", ABSTRACT_CN)
    abstract_en_count = replace_section_after_heading(doc, "ABSTRACT", ABSTRACT_EN)
    title_count = replace_title_text(doc)
    theory_count = replace_exact_paragraphs(doc, THEORY_REPLACEMENTS)
    parentheses_count = apply_parentheses_to_doc(doc)
    table_count = set_tables_for_cross_page(doc)

    doc.save(DOC_PATH)
    print(f"updated={DOC_PATH}")
    print(f"backup={backup}")
    print(
        "changed "
        f"abstract_cn={abstract_cn_count} abstract_en={abstract_en_count} "
        f"titles={title_count} theory={theory_count} "
        f"parentheses={parentheses_count} repeated_headers={table_count}"
    )


if __name__ == "__main__":
    main()
