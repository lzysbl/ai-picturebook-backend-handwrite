"""论文 DOCX 构建脚本。

用途：
- 将论文正文 Markdown/文本内容写入学校模板 DOCX。
- 处理封面、摘要、目录、正文标题、图片、表格、页眉页脚和页码等格式。
- 用于生成可提交或继续人工校对的论文 Word 文档。

关联文件：
- 读取报告目录中的论文源文本、图片和学校模板。
- 输出最终论文 DOCX。

运行方式：
- `python scripts/build_thesis_docx.py`
"""

from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path.cwd()
SOURCE_NAME = "毕业论文初稿.md"
OUTPUT_NAME = "20225973_姚家路_毕业论文_附件9套版.docx"
ENGLISH_TITLE = "Design and Implementation of a Picture-Book Storytelling Application Based on Multimodal Large Language Models"


def find_source_markdown() -> Path:
    for path in ROOT.rglob("*.md"):
        if path.name == SOURCE_NAME:
            return path
    raise FileNotFoundError(f"Could not find {SOURCE_NAME}")


def output_path() -> Path:
    for path in ROOT.iterdir():
        if path.is_dir() and path.name == "报告":
            return path / OUTPUT_NAME
    report_dir = ROOT / "报告"
    report_dir.mkdir(parents=True, exist_ok=True)
    return report_dir / OUTPUT_NAME


def find_attachment9_template() -> Path:
    for path in ROOT.rglob("*.docx"):
        if path.name.startswith("~$"):
            continue
        if "附件9" in path.name and "书写印制规范" in path.name:
            return path
    # Fallback: the official writing specification/template is the largest docx in the report folder.
    docx_files = [path for path in ROOT.rglob("*.docx") if not path.name.startswith("~$")]
    if not docx_files:
        raise FileNotFoundError("Could not find any docx template")
    return max(docx_files, key=lambda p: p.stat().st_size)


def capture_attachment9_sections(doc: Document) -> dict[str, object]:
    """Copy native section properties from the example area of attachment 9.

    The first two section breaks in the file belong to the writing specification
    pages. The later breaks are the actual example thesis structure: Chinese
    cover, English cover, declaration, abstracts, TOC, body, and references.
    """
    section_breaks = []
    for paragraph in doc.paragraphs:
        ppr = paragraph._p.pPr
        if ppr is not None and ppr.sectPr is not None:
            section_breaks.append(deepcopy(ppr.sectPr))
    if len(section_breaks) < 9:
        raise RuntimeError("Attachment 9 template does not contain enough section breaks")
    return {
        "cover": section_breaks[2],
        "english_cover": section_breaks[3],
        "declaration": section_breaks[4],
        "abstract_zh": section_breaks[5],
        "abstract_en": section_breaks[6],
        "toc": section_breaks[7],
        "body": section_breaks[8],
    }


def clear_document_body(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def append_template_section_break(doc: Document, sect_pr) -> None:
    paragraph = doc.add_paragraph()
    ppr = paragraph._p.get_or_add_pPr()
    existing = ppr.find(qn("w:sectPr"))
    if existing is not None:
        ppr.remove(existing)
    ppr.append(deepcopy(sect_pr))


def save_document(doc: Document, preferred: Path) -> Path:
    try:
        doc.save(str(preferred))
        return preferred
    except PermissionError:
        stem = preferred.stem
        suffix = preferred.suffix
        for index in range(2, 100):
            candidate = preferred.with_name(f"{stem}_{index}{suffix}")
            try:
                doc.save(str(candidate))
                return candidate
            except PermissionError:
                continue
        raise


def set_rfonts(run, east: str = "宋体", west: str = "Times New Roman"):
    run.font.name = west
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east)
    run._element.rPr.rFonts.set(qn("w:ascii"), west)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), west)


def set_run(run, *, east: str = "宋体", west: str = "Times New Roman", size: float = 12, bold: bool | None = None):
    set_rfonts(run, east=east, west=west)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if bold is not None:
        run.bold = bold


def para_fmt(paragraph, *, first_line: bool = False, line_pt: float | None = 23, before: float = 0, after: float = 0, align=None):
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Pt(24) if first_line else Pt(0)
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    if line_pt is None:
        fmt.line_spacing_rule = None
        fmt.line_spacing = None
    else:
        fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        fmt.line_spacing = Pt(line_pt)
    if align is not None:
        paragraph.alignment = align


def add_text(
    doc: Document,
    text: str,
    *,
    east: str = "宋体",
    west: str = "Times New Roman",
    size: float = 12,
    bold: bool | None = None,
    first_line: bool = False,
    line_pt: float | None = 23,
    before: float = 0,
    after: float = 0,
    align=None,
):
    p = doc.add_paragraph()
    para_fmt(p, first_line=first_line, line_pt=line_pt, before=before, after=after, align=align)
    r = p.add_run(text)
    set_run(r, east=east, west=west, size=size, bold=bold)
    return p


def add_field(paragraph, code: str):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = code
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(end)
    return run


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    left = paragraph.add_run("-")
    set_run(left, east="宋体", size=10.5)
    add_field(paragraph, "PAGE")
    right = paragraph.add_run("-")
    set_run(right, east="宋体", size=10.5)


def set_paragraph_bottom_border(paragraph, *, color: str = "D9D9D9", size: str = "4", space: str = "1"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:color"), color)
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)


def set_layout(section):
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    section.gutter = Cm(0)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)


def add_header_footer(section, right_title: str | None, *, show_header: bool = True, show_footer: bool = True):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    for p in section.header.paragraphs:
        p.text = ""
    for p in section.footer.paragraphs:
        p.text = ""

    if show_header:
        p = section.header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        tab_pos = section.page_width - section.left_margin - section.right_margin
        p.paragraph_format.tab_stops.add_tab_stop(tab_pos, WD_TAB_ALIGNMENT.RIGHT)
        r1 = p.add_run("东北大学本科生毕业设计（论文）")
        set_run(r1, east="宋体", size=10.5)
        if right_title:
            p.add_run("\t")
            r2 = p.add_run(right_title)
            set_run(r2, east="宋体", size=10.5)
        set_paragraph_bottom_border(p)
    if show_footer:
        p = section.footer.paragraphs[0]
        add_page_number(p)
        for r in p.runs:
            set_run(r, east="宋体", size=10.5)


def set_page_numbering(section, fmt: str | None = None, start: int | None = None):
    sectPr = section._sectPr
    pg = sectPr.find(qn("w:pgNumType"))
    if pg is None:
        pg = OxmlElement("w:pgNumType")
        sectPr.append(pg)
    if fmt is not None:
        pg.set(qn("w:fmt"), fmt)
    if start is not None:
        pg.set(qn("w:start"), str(start))


def new_section(
    doc: Document,
    right_title: str | None,
    *,
    show_header: bool = True,
    show_footer: bool = True,
    start_type=WD_SECTION_START.ODD_PAGE,
    page_fmt: str | None = None,
    page_start: int | None = None,
):
    section = doc.add_section(start_type)
    set_layout(section)
    add_header_footer(section, right_title, show_header=show_header, show_footer=show_footer)
    set_page_numbering(section, fmt=page_fmt, start=page_start)
    return section


def configure_doc(doc: Document):
    set_layout(doc.sections[0])
    add_header_footer(doc.sections[0], None, show_header=False, show_footer=False)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    doc.settings.element.append(update_fields)


def extract_meta(lines: list[str]) -> dict[str, str]:
    meta: dict[str, str] = {}
    for line in lines[:20]:
        m = re.match(r"^([^：:]+)[：:]\s*(.+)$", line.strip())
        if m:
            meta[m.group(1)] = m.group(2)
    return meta


def split_between(text: str, start: str, end: str | None) -> str:
    lines = text.splitlines()
    start_idx = None
    end_idx = len(lines)
    for i, line in enumerate(lines):
        if start_idx is None and line.startswith(start):
            start_idx = i + 1
            continue
        if start_idx is not None and end is not None and line.startswith(end):
            end_idx = i
            break
    if start_idx is None:
        return ""
    return "\n".join(lines[start_idx:end_idx]).strip()


def body_after_toc(text: str) -> str:
    lines = text.splitlines()
    for i, line in enumerate(lines):
        if line.startswith("## 目录"):
            for j in range(i + 1, len(lines)):
                if lines[j].startswith("# "):
                    return "\n".join(lines[j:]).strip()
            return ""
    return text


def add_cover(doc: Document, meta: dict[str, str], title: str):
    add_text(
        doc,
        f"学号    {meta.get('学号', '')}                                    密级________________",
        east="黑体",
        size=10.5,
        line_pt=23,
        first_line=True,
        align=WD_ALIGN_PARAGRAPH.LEFT,
    )
    add_text(doc, "", line_pt=23)
    add_text(doc, "东北大学本科生毕业设计（论文）", east="宋体", size=26, bold=True, line_pt=32, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "", line_pt=23)
    add_text(doc, title, east="黑体", size=22, bold=True, line_pt=32, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "", line_pt=23)
    add_text(doc, "", line_pt=23)
    add_text(doc, f"学 院 名 称  ：{meta.get('学院', '')}", east="宋体", size=15, line_pt=28, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, f"专 业 名 称  ：{meta.get('专业', '')}", east="宋体", size=15, line_pt=28, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, f"学 生 姓 名  ：{meta.get('姓名', '')}", east="宋体", size=15, line_pt=28, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, f"指 导 教 师  ：{meta.get('指导教师', '')}", east="宋体", size=15, line_pt=28, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "", line_pt=23)
    add_text(doc, "", line_pt=23)
    add_text(doc, "东北大学", east="宋体", size=15, line_pt=28, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "2026年5月", east="宋体", size=16, line_pt=28, align=WD_ALIGN_PARAGRAPH.CENTER)


def add_english_cover(doc: Document):
    title_lines = [
        "Design and Implementation of a",
        "Picture-Book Storytelling Application",
        "Based on Multimodal Large Language Models",
    ]
    for _ in range(4):
        add_text(doc, "", line_pt=23)
    for line in title_lines:
        add_text(
            doc,
            line,
            east="Times New Roman",
            west="Times New Roman",
            size=18,
            bold=True,
            line_pt=28,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
    for _ in range(8):
        add_text(doc, "", line_pt=23)
    add_text(doc, "Northeastern University", east="Times New Roman", west="Times New Roman", size=15, line_pt=28, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "May 2026", east="Times New Roman", west="Times New Roman", size=15, line_pt=28, align=WD_ALIGN_PARAGRAPH.CENTER)


def add_declaration(doc: Document):
    add_text(doc, "郑 重 声 明", east="宋体", size=22, bold=True, line_pt=23, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(
        doc,
        "本人呈交的毕业设计（论文），是在导师的指导下，独立进行研究工作所取得的成果，所有数据、图片资料真实可靠。尽我所知，除文中已经注明引用的内容外，本毕业设计（论文）的研究成果不包含他人享有著作权的内容。对本论文所涉及的研究工作做出贡献的其他个人和集体，均已在文中以明确的方式标明。",
        east="宋体",
        size=14,
        first_line=True,
        line_pt=23,
        align=WD_ALIGN_PARAGRAPH.LEFT,
    )
    add_text(doc, "", line_pt=23)
    add_text(doc, "", line_pt=23)
    add_text(doc, "本人签名：                          日期：", east="宋体", size=14, line_pt=23, align=WD_ALIGN_PARAGRAPH.LEFT)


def add_abs_zh(doc: Document, block: str):
    add_text(doc, "摘  要", east="黑体", size=18, before=12.45, after=7.8, line_pt=None, align=WD_ALIGN_PARAGRAPH.CENTER)
    for para in block.splitlines():
        line = para.strip()
        if not line:
            continue
        if line.startswith("关键词"):
            add_text(doc, "", line_pt=23)
            p = doc.add_paragraph()
            para_fmt(p, first_line=False, line_pt=23)
            label, _, rest = line.partition("：")
            r1 = p.add_run(label + "：")
            set_run(r1, east="黑体", size=12, bold=False)
            r2 = p.add_run(rest)
            set_run(r2, east="宋体", size=12)
        else:
            add_text(doc, line, east="宋体", size=12, first_line=True, line_pt=23)


def add_abs_en(doc: Document, block: str):
    add_text(doc, "ABSTRACT", east="Times New Roman", west="Times New Roman", size=18, bold=True, before=9.6, after=6, line_pt=None, align=WD_ALIGN_PARAGRAPH.CENTER)
    for para in block.splitlines():
        line = para.strip()
        if not line:
            continue
        if line.startswith("Keywords"):
            add_text(doc, "", line_pt=23)
            p = doc.add_paragraph()
            para_fmt(p, first_line=False, line_pt=23)
            label, _, rest = line.partition(":")
            r1 = p.add_run(label + ":")
            set_run(r1, east="Times New Roman", west="Times New Roman", size=12, bold=True)
            r2 = p.add_run(rest)
            set_run(r2, east="Times New Roman", west="Times New Roman", size=12)
        else:
            add_text(doc, line, east="Times New Roman", west="Times New Roman", size=12, first_line=True, line_pt=23)


def add_toc(doc: Document):
    add_text(doc, "目  录", east="黑体", size=18, bold=True, before=9.6, after=6, line_pt=None, align=WD_ALIGN_PARAGRAPH.CENTER)
    p = doc.add_paragraph()
    para_fmt(p, first_line=False, line_pt=None)
    add_field(p, r'TOC \o "1-3" \h \z \u')


def add_heading(doc: Document, title: str, level: int):
    if level == 1:
        if len(doc.paragraphs) > 0:
            new_section(doc, title, page_fmt="decimal", page_start=1)
        p = add_text(doc, title, east="黑体", size=18, before=12.45, after=7.8, line_pt=None, align=WD_ALIGN_PARAGRAPH.CENTER)
        p.style = doc.styles["Heading 1"]
        return p
    if level == 2:
        p = add_text(doc, title, east="黑体", size=14, before=7.8, after=7.8, line_pt=None, align=WD_ALIGN_PARAGRAPH.LEFT)
        p.style = doc.styles["Heading 2"]
        return p
    p = add_text(doc, title, east="黑体", size=12, before=0, after=0, line_pt=23, align=WD_ALIGN_PARAGRAPH.LEFT)
    p.style = doc.styles["Heading 3"]
    return p


def add_table_from_rows(doc: Document, rows: list[str]):
    header = [c.strip() for c in rows[0].strip("|").split("|")]
    body_rows = [[c.strip() for c in row.strip("|").split("|")] for row in rows[2:]]
    table = doc.add_table(rows=1 + len(body_rows), cols=len(header))
    table.style = "Table Grid"
    for j, text in enumerate(header):
        table.rows[0].cells[j].text = text
    for i, row in enumerate(body_rows, 1):
        for j, text in enumerate(row[: len(header)]):
            table.rows[i].cells[j].text = text
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                para_fmt(p, first_line=False, line_pt=15, align=WD_ALIGN_PARAGRAPH.CENTER)
                for r in p.runs:
                    set_run(r, east="宋体", size=10.5)


def is_caption(line: str) -> bool:
    return bool(re.match(r"^(表|图)\s*\d+(?:\.\d+)?\s+", line))


def add_markdown_image(doc: Document, alt_text: str, image_ref: str):
    image_path = Path(image_ref)
    if not image_path.is_absolute():
        image_path = ROOT / image_path
    if not image_path.exists():
        add_text(doc, f"[缺少图片: {alt_text}] {image_ref}", east="宋体", size=10.5, line_pt=23, align=WD_ALIGN_PARAGRAPH.CENTER)
        return
    paragraph = doc.add_paragraph()
    para_fmt(paragraph, first_line=False, line_pt=None, before=7.8, after=7.8, align=WD_ALIGN_PARAGRAPH.CENTER)
    available_width = doc.sections[-1].page_width - doc.sections[-1].left_margin - doc.sections[-1].right_margin
    paragraph.add_run().add_picture(str(image_path), width=min(available_width, Cm(15.2)))
    if alt_text:
        caption = doc.add_paragraph()
        para_fmt(caption, first_line=False, line_pt=23, before=0, after=7.8, align=WD_ALIGN_PARAGRAPH.CENTER)
        run = caption.add_run(alt_text)
        set_run(run, east="宋体", size=10.5)


def add_body_block(doc: Document, block: str):
    lines = block.splitlines()
    first_heading = True
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        image_match = re.match(r"^!\[(.*?)\]\((.*?)\)$", line)
        if image_match:
            add_markdown_image(doc, image_match.group(1).strip(), image_match.group(2).strip())
            i += 1
            continue
        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            title = line[level:].strip()
            if first_heading and level == 1:
                add_text(doc, title, east="黑体", size=18, before=12.45, after=7.8, line_pt=None, align=WD_ALIGN_PARAGRAPH.CENTER).style = doc.styles["Heading 1"]
                first_heading = False
            else:
                if level == 1:
                    new_section(doc, title, page_fmt="decimal")
                    add_text(doc, title, east="黑体", size=18, before=12.45, after=7.8, line_pt=None, align=WD_ALIGN_PARAGRAPH.CENTER).style = doc.styles["Heading 1"]
                else:
                    add_heading(doc, title, min(level, 3))
            i += 1
            continue
        if line.startswith("|"):
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_rows.append(lines[i].strip())
                i += 1
            if len(table_rows) >= 3:
                add_table_from_rows(doc, table_rows)
            continue
        if line.startswith("- ") or line.startswith("* "):
            add_text(doc, line[2:].strip(), east="宋体", size=12, first_line=True, line_pt=23)
            i += 1
            continue
        if line.startswith("```"):
            i += 1
            code = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            add_text(doc, "\n".join(code), east="Courier New", west="Courier New", size=10, first_line=False, line_pt=15)
            if i < len(lines):
                i += 1
            continue
        if is_caption(line):
            add_text(doc, line, east="宋体", size=10.5, before=7.8, after=7.8, line_pt=23, align=WD_ALIGN_PARAGRAPH.CENTER)
        elif re.match(r"^\[\d+\]", line):
            add_text(doc, line, east="宋体", size=12, first_line=False, line_pt=23)
        else:
            add_text(doc, line, east="宋体", size=12, first_line=True, line_pt=23)
        i += 1


def build_document() -> Path:
    source = find_source_markdown()
    text = source.read_text(encoding="utf-8")
    lines = text.splitlines()
    title = lines[0].lstrip("# ").strip()
    meta = extract_meta(lines)
    zh_abs = split_between(text, "## 摘要", "## Abstract")
    en_abs = split_between(text, "## Abstract", "## 目录")
    body = body_after_toc(text)

    template_path = find_attachment9_template()
    template_doc = Document(str(template_path))
    section_breaks = capture_attachment9_sections(template_doc)
    doc = Document(str(template_path))
    clear_document_body(doc)
    configure_doc(doc)
    set_page_numbering(doc.sections[0])
    add_cover(doc, meta, title)
    append_template_section_break(doc, section_breaks["english_cover"])
    add_english_cover(doc)
    append_template_section_break(doc, section_breaks["declaration"])
    add_declaration(doc)
    append_template_section_break(doc, section_breaks["abstract_zh"])
    add_abs_zh(doc, zh_abs)
    append_template_section_break(doc, section_breaks["abstract_en"])
    add_abs_en(doc, en_abs)
    append_template_section_break(doc, section_breaks["toc"])
    add_toc(doc)
    append_template_section_break(doc, section_breaks["body"])
    add_header_footer(doc.sections[-1], "1 绪论")
    set_page_numbering(doc.sections[-1], fmt="decimal", start=1)
    add_body_block(doc, body)

    return save_document(doc, output_path())


def main() -> None:
    print(build_document())


if __name__ == "__main__":
    main()
