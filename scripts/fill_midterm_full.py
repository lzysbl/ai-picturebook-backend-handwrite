from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


def set_run_font(run, east="宋体", west="Times New Roman", size=12):
    run.font.name = west
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east)
    run.font.size = Pt(size)


def set_cell_text(cell, text, east="宋体", size=12):
    cell.text = text
    for para in cell.paragraphs:
        for run in para.runs:
            set_run_font(run, east=east, size=size)


def find_midterm_template() -> Path:
    # 优先找“报告/中期报告.docx”
    preferred = Path("报告") / "中期报告.docx"
    if preferred.exists():
        return preferred

    for p in Path(".").rglob("*.docx"):
        esc = p.name.encode("unicode_escape").decode("ascii")
        if "\\u4e2d\\u671f\\u62a5\\u544a" in esc:
            return p
    raise FileNotFoundError("未找到中期报告模板（*.docx）")


summary_text = (
    "前一阶段的工作围绕“基于多模态大模型的绘本讲述应用”后端系统原型展开，已完成需求拆解、技术路线确定、"
    "数据库建模与核心模块编码。首先，在文献调研方面，已系统梳理多模态大模型、视觉叙事、图文对齐与评估指标相关研究，"
    "明确系统需具备“多页绘本输入处理、语义结构化表示、讲述生成与质量评估”四个关键能力。其次，在系统设计方面，"
    "已采用分层架构组织工程，目录划分为 core、db、models、schemas、services、routers、utils，"
    "确保配置、数据访问、业务逻辑和接口层解耦，便于后续扩展与维护。在数据层，已完成 users、books、book_images、stories "
+    "四张核心表及关联关系设计，实现“用户-绘本-图片-故事”的完整链路。在实现层，已完成服务层关键模块的异步化重构，"
    "统一为 FastAPI + SQLAlchemy 异步风格，去除了与项目技术栈不一致的写法；并补充密码哈希和 JWT 令牌生成能力，"
    "为注册登录接口联调提供基础。在工具层，已完成绘本图片爬虫的并发下载改造和保存目录校验，提升样本数据准备效率。"
    "在文档层，README 已更新为项目化描述，补充 Schemas 图与数据库 ER 图，为中期汇报与论文撰写提供结构化材料。"
    "总体上，当前阶段已经完成从“概念与调研”向“可运行工程骨架”的转化，为下一阶段功能闭环、实验验证和论文实现章节奠定了基础。"
)

plan_text = (
    "下一阶段将以“完成可运行闭环+形成论文实验支撑”为目标推进。第一，完成用户模块的注册、登录、鉴权联调，"
    "将 JWT 鉴权接入受保护接口，规范错误码与异常返回。第二，完成绘本与图片模块端到端调试，覆盖绘本创建、图片上传、"
    "顺序管理、文件落盘与数据库记录一致性校验。第三，完成故事模块闭环，实现图片分析、故事生成、故事入库、历史查询、详情查询，"
    "并对关键流程增加输入参数校验。第四，完善 AI 服务层接口形态，保留 mock 逻辑的同时统一抽象，便于后续替换真实多模态模型 API。"
    "第五，补充质量评估输出，围绕图文对应性、叙事连贯性与适龄性形成可解释评分，支持结果展示。第六，完善工程文档与测试，"
    "包括 API 调试示例、运行说明、健康检查与关键服务函数测试用例。论文方面，将同步整理系统分析、系统设计与实现章节素材，"
    "沉淀图表、接口截图与阶段实验记录，保证系统实现与论文内容一致，按节点推进中后期答辩准备。"
)

problem_text = (
    "当前主要问题包括：一是异步数据库环境依赖（如驱动版本）对联调稳定性有影响，环境不一致时容易出现启动和导入报错；"
    "二是部分模块虽然已完成代码重构，但接口联调尚未全部闭环，尤其在异常处理、参数校验和统一响应方面仍需细化；"
    "三是 AI 讲述能力目前为 mock，实现了流程验证但尚未接入真实多模态模型，后续可能在时延、稳定性和输出质量上出现新问题。"
    "针对上述问题，后续将通过固定环境版本、补齐测试、分阶段联调和指标化评估降低风险。"
)

other_text = (
    "开题报告参考文献要点摘要：\n"
    "[1] Yin 等（2023）综述多模态大模型总体框架、训练策略与评估方法，指出视觉叙事与多模态推理潜力。\n"
    "[2] Yin 等（2024）从模型构成与发展趋势角度总结 MLLM，强调跨模态理解与生成一体化能力。\n"
    "[3] Liang 等（2024）系统梳理视觉-语言任务中的 MLLM 实践路径，为工程落地提供方法参考。\n"
    "[4] Han 等（2025）进一步总结新一代 MLLM，提出一致性、可控性与泛化能力仍是关键挑战。\n"
    "[5] Gado 等（2025）聚焦视觉故事生成，讨论 VIST-GPT 及相关评估问题，强调叙事质量指标的重要性。\n"
    "[6] Yang 等（2025，StoryLLaVA）提出话题驱动和偏好策略，提升视觉故事连贯性与趣味性。\n"
    "[7] Bai 等（2024）综述多模态幻觉问题，提示在生成系统中需加入一致性约束与校验机制。\n"
    "[8] 张沁彦（2024）从多模态语篇角度分析儿童绘本意义建构，对本课题讲述策略有参考价值。\n"
    "[9] 冯海林课题组（2024）讨论基础模型向多模态模型演进，为技术路线选型提供理论支持。\n"
    "[10] 53AI（2024）提供多模态大模型中文综述资料，便于对国内技术生态进行补充调研。"
)

teacher_opinion_text = (
    "指导教师意见（建议稿）：\n"
    "该生前期工作推进较为扎实，已完成系统架构划分、数据库建模与核心服务层重构，能够围绕课题目标开展工程实现。"
    "中期阶段已形成可持续迭代的后端原型，技术路线与研究内容匹配度较高。建议下一阶段重点完成接口联调闭环、"
    "真实模型接入与评估指标实验，并加强测试与论文材料同步整理。总体进度符合中期要求。\n"
    "（注：此处最终内容请由指导教师确认并签字）"
)


def main():
    src = find_midterm_template()
    doc = Document(str(src))

    # 日期行
    if len(doc.paragraphs) >= 2:
        doc.paragraphs[1].text = "                                                               2026年4月10日"
        for run in doc.paragraphs[1].runs:
            set_run_font(run)

    table = doc.tables[0]
    # 约定：第4-8行分别是 前一阶段总结/下一阶段计划/存在问题/其它/指导教师意见
    set_cell_text(table.rows[3].cells[1], summary_text)
    set_cell_text(table.rows[4].cells[1], plan_text)
    set_cell_text(table.rows[5].cells[1], problem_text)
    set_cell_text(table.rows[6].cells[1], other_text)
    set_cell_text(table.rows[7].cells[1], teacher_opinion_text)

    # 全局字体兜底
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    out_dir = Path("报告")
    out_dir.mkdir(parents=True, exist_ok=True)
    out = out_dir / "midterm_report_full.docx"
    doc.save(str(out))
    print(out.resolve())


if __name__ == "__main__":
    main()
