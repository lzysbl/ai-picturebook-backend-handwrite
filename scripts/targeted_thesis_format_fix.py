from __future__ import annotations

import re
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


REPORT_DIR = Path("报告")
DOC_PATH = REPORT_DIR / "论文第二版.docx"
BACKUP_PATH = REPORT_DIR / f"论文第二版_backup_before_targeted_format_{datetime.now():%Y%m%d_%H%M%S}.docx"

BODY_EAST = "宋体"
HEADING_EAST = "黑体"
LATIN = "Times New Roman"
SOFTWARE_FIG_DIR = Path("reports") / "software_engineering_figures"
THESIS_FIG_DIR = Path("reports") / "thesis_figures"

FIGURE_SPECS = [
    {
        "caption": "图3.1 系统用例图",
        "path": SOFTWARE_FIG_DIR / "figure_3_1_use_case.png",
        "anchor": "在功能需求基础上，系统用例可以抽象为用户与外部服务之间的交互关系。",
        "width": 5.85,
    },
    {
        "caption": "图4.1 系统总体架构",
        "path": THESIS_FIG_DIR / "figure_4_1_system_architecture.png",
        "anchor": "总体上，前端是静态 HTML、CSS 和 JavaScript",
        "width": 5.95,
    },
    {
        "caption": "图4.2 数据库实体关系图",
        "path": SOFTWARE_FIG_DIR / "figure_4_4_database_er.png",
        "anchor": "各数据表之间的关系如图4.2 所示。",
        "width": 5.95,
    },
    {
        "caption": "图4.3 实时绘本识别与讲述流程",
        "path": THESIS_FIG_DIR / "figure_4_2_live_scan_pipeline.png",
        "anchor": "实时识别模块主要支持以下能力。",
        "width": 5.95,
    },
    {
        "caption": "图4.4 后端功能模块划分",
        "path": THESIS_FIG_DIR / "figure_4_3_backend_modules.png",
        "anchor": "图4.3 实时绘本识别与讲述流程",
        "width": 5.95,
    },
    {
        "caption": "图5.1 实时识别与朗读顺序图",
        "path": SOFTWARE_FIG_DIR / "figure_5_3_live_scan_sequence.png",
        "anchor": "本文将实时识别与朗读过程整理为图5.1。",
        "width": 6.05,
    },
    {
        "caption": "图5.2 前端实时识别状态流转",
        "path": THESIS_FIG_DIR / "figure_5_1_frontend_state.png",
        "anchor": "前端页面也围绕手机端做了调整。",
        "width": 5.95,
    },
    {
        "caption": "图5.3 语音朗读处理流程",
        "path": THESIS_FIG_DIR / "figure_5_2_tts_flow.png",
        "anchor": "生成阶段还加入了角色一致性约束。",
        "width": 5.85,
    },
    {
        "caption": "图6.1 不同识别与讲述方案延迟对比",
        "path": THESIS_FIG_DIR / "figure_6_1_latency_comparison.png",
        "anchor": "图6.1 不同识别与讲述方案延迟对比",
        "width": 5.85,
    },
    {
        "caption": "图6.2 运行日志主要耗时指标",
        "path": THESIS_FIG_DIR / "figure_6_2_runtime_metrics.png",
        "anchor": "图6.2 运行日志主要耗时指标",
        "width": 5.85,
    },
    {
        "caption": "图6.3 测试数据来源构成",
        "path": THESIS_FIG_DIR / "figure_6_3_dataset_composition.png",
        "anchor": "图6.3 测试数据来源构成",
        "width": 5.85,
    },
]


def clean(text: str) -> str:
    return " ".join((text or "").replace("\u200c", "").split())


def ensure_rpr(run):
    rpr = run._element.get_or_add_rPr()
    if rpr.rFonts is None:
        rpr.append(OxmlElement("w:rFonts"))
    return rpr


def set_run_font(
    run,
    *,
    east: str = BODY_EAST,
    latin: str = LATIN,
    size_pt: float | None = 12,
    bold: bool | None = None,
    superscript: bool | None = None,
) -> None:
    run.font.name = latin
    rpr = ensure_rpr(run)
    rfonts = rpr.rFonts
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if superscript is not None:
        run.font.superscript = superscript


def set_para_line(paragraph, *, first_indent: bool = False, center: bool = False) -> None:
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(22)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.first_line_indent = Pt(24) if first_indent else Pt(0)
    if center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def is_drawing_paragraph(element) -> bool:
    return element.tag == qn("w:p") and bool(element.xpath(".//w:drawing"))


def has_visible_text(element) -> bool:
    if element.tag != qn("w:p"):
        return False
    text = "".join(t.text or "" for t in element.iter(qn("w:t")))
    return bool(clean(text))


def clear_runs(paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def apply_superscript_citations(paragraph) -> None:
    text = paragraph.text
    if not re.search(r"\[\d+(?:\]\[\d+)*\]", text):
        return
    parts = re.split(r"(\[\d+(?:\]\[\d+)*\])", text)
    clear_runs(paragraph)
    for part in parts:
        if part == "":
            continue
        run = paragraph.add_run(part)
        is_cite = bool(re.fullmatch(r"\[\d+(?:\]\[\d+)*\]", part))
        set_run_font(run, superscript=is_cite)


def set_cell_text(cell, value: str, bold: bool = False) -> None:
    cell.text = value
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(18)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            set_run_font(run, size_pt=9, bold=bold)


def border_element(name: str, val: str = "nil", size: str = "0"):
    element = OxmlElement(f"w:{name}")
    element.set(qn("w:val"), val)
    element.set(qn("w:sz"), size)
    element.set(qn("w:space"), "0")
    element.set(qn("w:color"), "000000" if val != "nil" else "auto")
    return element


def set_cell_border(cell, top: str | None = None, bottom: str | None = None) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for old in tc_pr.findall(qn("w:tcBorders")):
        tc_pr.remove(old)
    borders = OxmlElement("w:tcBorders")
    borders.append(border_element("top", "single" if top else "nil", top or "0"))
    borders.append(border_element("left"))
    borders.append(border_element("bottom", "single" if bottom else "nil", bottom or "0"))
    borders.append(border_element("right"))
    borders.append(border_element("insideH"))
    borders.append(border_element("insideV"))
    tc_pr.append(borders)


def make_three_line(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    tbl_pr = table._tbl.tblPr
    for old in tbl_pr.findall(qn("w:tblBorders")):
        tbl_pr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        borders.append(border_element(name))
    tbl_pr.append(borders)
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell)
    if len(table.rows) > 0:
        for cell in table.rows[0].cells:
            set_cell_border(cell, top="12", bottom="8")
        for cell in table.rows[-1].cells:
            set_cell_border(cell, bottom="12")


def find_caption_index(doc: Document, caption_text: str) -> int:
    target = clean(caption_text)
    for i, p in enumerate(doc.paragraphs):
        if clean(p.text) == target:
            return i
    raise ValueError(f"caption not found: {caption_text}")


def find_paragraph_index_containing(doc: Document, text: str) -> int:
    for i, p in enumerate(doc.paragraphs):
        if text in clean(p.text):
            return i
    raise ValueError(f"anchor paragraph not found: {text}")


def insert_paragraph_after(anchor_element, paragraph_element) -> None:
    parent = paragraph_element.getparent()
    if parent is not None:
        parent.remove(paragraph_element)
    anchor_element.addnext(paragraph_element)


def insert_paragraph_before(anchor_element, paragraph_element) -> None:
    parent = paragraph_element.getparent()
    if parent is not None:
        parent.remove(paragraph_element)
    anchor_element.addprevious(paragraph_element)


def table_after_caption(doc: Document, caption_text: str):
    body_children = list(doc.element.body)
    caption_p = doc.paragraphs[find_caption_index(doc, caption_text)]._p
    start = body_children.index(caption_p)
    for element in body_children[start + 1 :]:
        if element.tag == qn("w:tbl"):
            for table in doc.tables:
                if table._tbl is element:
                    return table
        if element.tag == qn("w:p"):
            text = clean("".join(t.text or "" for t in element.iter(qn("w:t"))))
            if text.startswith(("表 ", "图", "表")) and text != clean(caption_text):
                break
    raise ValueError(f"table after caption not found: {caption_text}")


def add_caption_paragraph(doc: Document, caption_text: str):
    p = doc.add_paragraph(caption_text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_line(p, center=True)
    for run in p.runs:
        set_run_font(run, size_pt=10.5)
    return p


def add_picture_paragraph(doc: Document, image_path: Path, width_inches: float):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Do not use fixed 22 pt line spacing on picture-only paragraphs.
    # Word clips inline pictures when the containing paragraph has exact line
    # spacing, which makes large figures look missing even though they are
    # embedded in the DOCX package.
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(0)
    p.add_run().add_picture(str(image_path), width=Inches(width_inches))
    return p


def ensure_figure(doc: Document, spec: dict[str, object]) -> None:
    caption_text = str(spec["caption"])
    image_path = Path(spec["path"])
    anchor_text = str(spec["anchor"])
    width = float(spec["width"])

    if not image_path.exists():
        raise FileNotFoundError(image_path)

    body_children = list(doc.element.body)
    try:
        caption_index = find_caption_index(doc, caption_text)
        caption_p = doc.paragraphs[caption_index]
    except ValueError:
        anchor_index = find_paragraph_index_containing(doc, anchor_text)
        anchor_p = doc.paragraphs[anchor_index]
        picture_p = add_picture_paragraph(doc, image_path, width)
        caption_p = add_caption_paragraph(doc, caption_text)
        insert_paragraph_after(anchor_p._p, caption_p._p)
        insert_paragraph_after(anchor_p._p, picture_p._p)
        return

    caption_el = caption_p._p
    caption_pos = body_children.index(caption_el)
    previous_el = body_children[caption_pos - 1] if caption_pos > 0 else None
    if previous_el is not None and is_drawing_paragraph(previous_el) and not has_visible_text(previous_el):
        previous_el.getparent().remove(previous_el)

    picture_p = add_picture_paragraph(doc, image_path, width)
    insert_paragraph_before(caption_el, picture_p._p)


def ensure_figures(doc: Document) -> None:
    for spec in FIGURE_SPECS:
        ensure_figure(doc, spec)


def fill_table(table, rows: list[list[str]]) -> None:
    while len(table.rows) < len(rows):
        table.add_row()
    while len(table.rows) > len(rows):
        table._tbl.remove(table.rows[-1]._tr)
    for r, row_values in enumerate(rows):
        for c, value in enumerate(row_values):
            set_cell_text(table.rows[r].cells[c], value, bold=(r == 0))
    make_three_line(table)


def replace_text_in_paragraphs(doc: Document) -> None:
    replacements = {
        "/api/images/upload": "/api/books/{book_id}/images/upload",
        "/api/images/{book_id}/images/upload": "/api/books/{book_id}/images/upload",
        "绘本表存标题、描述和所属用户": "绘本表存标题、封面路径和所属用户",
        "可以进一步抽象为用户与外部服务之间的交互关系": "可以抽象为用户与外部服务之间的交互关系",
        "便于说明后端实现不是简单页面调用": "用于说明后端接口和前端页面之间的字段对应关系",
    }
    for p in doc.paragraphs:
        text = p.text
        new_text = text
        for src, dst in replacements.items():
            new_text = new_text.replace(src, dst)
        if new_text != text:
            clear_runs(p)
            run = p.add_run(new_text)
            set_run_font(run)


def replace_text_in_tables(doc: Document) -> None:
    replacements = {
        "/api/images/upload": "/api/books/{book_id}/images/upload",
        "/api/images/{book_id}/images/upload": "/api/books/{book_id}/images/upload",
        "username_or_email": "username",
        "image_order": "start_order",
        "style, target_age": "narration_style, audience_age",
        "mode, crop": "response_mode, crop_source/crop_box",
        "voice, rate, volume": "voice_preset",
        "current_text, total_story, analysis, timing": "analysis_result, story_content, quality, timing, context",
        "delta, final, timing": "meta/delta/done 事件, story_content, timing",
        "session_id, title, story_content, images": "story_content, page_stories, analysis_result, image_paths, session_id, book_id",
        "story_id, saved_at": "story, quality, book_id, image_paths",
    }
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                original = cell.text
                updated = original
                for src, dst in replacements.items():
                    updated = updated.replace(src, dst)
                if updated != original:
                    set_cell_text(cell, updated)


def update_interface_field_table(doc: Document) -> None:
    table = table_after_caption(doc, "表 5.2 核心接口请求与响应字段")
    rows = [
        ["接口", "主要请求字段", "主要响应字段", "说明"],
        ["/api/users/login", "username, password", "access_token, token_type, user", "完成用户认证并返回访问令牌"],
        ["/api/books", "title, cover_image", "id, user_id, title, cover_image, created_at", "创建或查询当前用户绘本"],
        [
            "/api/books/{book_id}/images/upload",
            "files, start_order",
            "id, book_id, image_path, image_order, created_at",
            "上传并保存绘本页面图片",
        ],
        [
            "/api/stories/generate",
            "book_id, prompt, narration_style, audience_age, story_length, generation_mode",
            "analysis_result, story_content, quality, story",
            "根据整本绘本生成故事",
        ],
        [
            "/api/stories/scan",
            "image, session_id, response_mode, crop_source/crop_box, narration_style, audience_age",
            "analysis_result, story_content, quality, timing, context",
            "识别当前页并返回讲述结果",
        ],
        [
            "/api/stories/scan/stream",
            "image, session_id, response_mode, crop_source/crop_box, narration_style, audience_age",
            "meta/delta/done 事件, story_content, timing",
            "以 SSE 方式流式返回讲述文本",
        ],
        ["/api/stories/tts", "text, voice_preset", "audio_url, provider, duration_seconds, timing", "将讲述文本转为语音文件"],
        [
            "/api/stories/scan/save",
            "story_content, page_stories, analysis_result, image_paths, session_id, book_id",
            "story, quality, book_id, image_paths",
            "保存实时扫描生成的总故事",
        ],
    ]
    fill_table(table, rows)


def update_core_api_table(doc: Document) -> None:
    table = table_after_caption(doc, "表 5.1 核心接口说明")
    for row in table.rows:
        if clean(row.cells[0].text) in {"`/api/images/upload`", "/api/images/upload"}:
            set_cell_text(row.cells[0], "`/api/books/{book_id}/images/upload`")
    make_three_line(table)


def format_document(doc: Document) -> None:
    for table in doc.tables[2:]:
        make_three_line(table)
        for r, row in enumerate(table.rows):
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                    p.paragraph_format.line_spacing = Pt(18)
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    for run in p.runs:
                        set_run_font(run, size_pt=9, bold=(r == 0 if run.text.strip() else None))

    in_references = False
    body_started = False
    for p in doc.paragraphs:
        text = clean(p.text)
        if p._p.xpath(".//w:drawing"):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.first_line_indent = Pt(0)
            continue
        if text == "摘 要":
            body_started = True
        if text == "参考文献":
            in_references = True
        style_name = p.style.name if p.style else ""

        if not text or style_name.lower().startswith("toc"):
            continue

        if style_name in {"Heading 1", "Heading 2"}:
            set_para_line(p)
            for run in p.runs:
                set_run_font(run, east=HEADING_EAST, size_pt=None, bold=True)
            continue

        is_caption = bool(re.match(r"^(图\s*\d|图\d|表\s*\d|表\d)", text))
        if is_caption:
            set_para_line(p, center=True)
            for run in p.runs:
                set_run_font(run, size_pt=10.5)
            continue

        if body_started and not in_references and style_name in {"Normal", "List Paragraph", "正文", "Body Text"}:
            if not text.startswith(("关键词", "Keywords", "`")):
                set_para_line(p, first_indent=True)
            else:
                set_para_line(p)
            apply_superscript_citations(p)
            for run in p.runs:
                superscript = bool(run.font.superscript)
                set_run_font(run, size_pt=12, superscript=superscript)
        elif in_references:
            set_para_line(p)
            for run in p.runs:
                set_run_font(run, size_pt=12)


def main() -> None:
    if not DOC_PATH.exists():
        raise FileNotFoundError(DOC_PATH)

    shutil.copy2(DOC_PATH, BACKUP_PATH)
    doc = Document(DOC_PATH)
    replace_text_in_paragraphs(doc)
    replace_text_in_tables(doc)
    update_core_api_table(doc)
    update_interface_field_table(doc)
    ensure_figures(doc)
    format_document(doc)
    doc.save(DOC_PATH)
    print(f"updated={DOC_PATH}")
    print(f"backup={BACKUP_PATH}")


if __name__ == "__main__":
    main()
