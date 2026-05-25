from __future__ import annotations

from datetime import datetime
from pathlib import Path
import shutil

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


REPORT_DIR = Path("报告")
DOC_PATH = REPORT_DIR / "论文第二版.docx"
FIG_DIR = Path("reports") / "software_engineering_figures"


def find_paragraph(doc: Document, text: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().replace("\u200c", "") == text:
            return paragraph
    raise ValueError(f"paragraph not found: {text}")


def find_following_table(doc: Document, caption: str):
    children = list(doc.element.body)
    caption_el = find_paragraph(doc, caption)._p
    start = children.index(caption_el)
    for element in children[start + 1 :]:
        if element.tag == qn("w:tbl"):
            return element
        if element.tag == qn("w:p"):
            text = "".join(t.text or "" for t in element.iter(qn("w:t"))).strip()
            if text.startswith(("表 ", "图")):
                break
    raise ValueError(f"table after caption not found: {caption}")


def move_after(anchor_el, elements):
    current = anchor_el
    for element in elements:
        if element.getparent() is not None:
            element.getparent().remove(element)
        current.addnext(element)
        current = element
    return current


def paragraph(doc: Document, text: str = "", style: str | None = None, center: bool = False):
    p = doc.add_paragraph(text)
    if style:
        p.style = style
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def caption(doc: Document, text: str):
    p = paragraph(doc, text, center=True)
    for run in p.runs:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(10.5)
    return p


def image_paragraph(doc: Document, path: Path, width=Inches(5.85)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=width)
    return p


def set_cell_text(cell, value: str, bold: bool = False):
    cell.text = value
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            run.font.size = Pt(9)
            run.bold = bold


def border_element(name: str, val: str = "nil", size: str = "0"):
    element = OxmlElement(f"w:{name}")
    element.set(qn("w:val"), val)
    element.set(qn("w:sz"), size)
    element.set(qn("w:space"), "0")
    element.set(qn("w:color"), "000000" if val != "nil" else "auto")
    return element


def set_cell_border(cell, top: str | None = None, bottom: str | None = None):
    tc_pr = cell._tc.get_or_add_tcPr()
    for old in tc_pr.findall(qn("w:tcBorders")):
        tc_pr.remove(old)
    borders = OxmlElement("w:tcBorders")
    borders.append(border_element("top", "single" if top else "nil", top or "0"))
    borders.append(border_element("left"))
    borders.append(border_element("bottom", "single" if bottom else "nil", bottom or "0"))
    borders.append(border_element("right"))
    borders.append(border_element("insideH"))
    borders.append(border_element("insideV"))
    tc_pr.append(borders)


def make_three_line(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    tbl_pr = table._tbl.tblPr
    for old in tbl_pr.findall(qn("w:tblBorders")):
        tbl_pr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        borders.append(border_element(name))
    tbl_pr.append(borders)
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell)
    for cell in table.rows[0].cells:
        set_cell_border(cell, top="12", bottom="8")
    for cell in table.rows[-1].cells:
        set_cell_border(cell, bottom="12")


def table(doc: Document, headers: list[str], rows: list[list[str]]):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for index, header in enumerate(headers):
        set_cell_text(t.rows[0].cells[index], header, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)
    make_three_line(t)
    return t


def replace_text(doc: Document):
    replacements = {
        "图4.2 实时绘本识别与讲述流程": "图4.3 实时绘本识别与讲述流程",
        "图4.3 后端功能模块划分": "图4.4 后端功能模块划分",
        "图5.1 前端实时识别状态流转": "图5.2 前端实时识别状态流转",
        "图5.2 语音朗读处理流程": "图5.3 语音朗读处理流程",
        "表 4.2 响应模式设计对比": "表 4.3 响应模式设计对比",
        "表 5.2 模型调用与提示词配置说明": "表 5.3 模型调用与提示词配置说明",
        "表 6.3 绘本识别延迟对比": "表 6.4 绘本识别延迟对比",
        "表 6.4 运行日志导出指标摘要": "表 6.5 运行日志导出指标摘要",
        "表 6.5 绘本讲述质量人工评价指标": "表 6.6 绘本讲述质量人工评价指标",
        "表 6.6 绘本讲述质量人工评分结果": "表 6.7 绘本讲述质量人工评分结果",
        "表 6.7 典型问题类型与改进方向": "表 6.8 典型问题类型与改进方向",
        "评价指标见表 6.5": "评价指标见表 6.6",
        "人工评分结果见表 6.6": "人工评分结果见表 6.7",
        "表 6.3 的批量 benchmark": "表 6.4 的批量 benchmark",
    }
    for p in doc.paragraphs:
        original = p.text
        if original.startswith("表 6.4 为开发与联调阶段运行日志的补充统计"):
            p.text = original.replace("表 6.4 为开发与联调阶段运行日志的补充统计", "运行日志补充统计")
            continue
        updated = original
        for source, target in replacements.items():
            updated = updated.replace(source, target)
        if updated != original:
            p.text = updated


def main():
    if not DOC_PATH.exists():
        raise FileNotFoundError(DOC_PATH)
    backup = REPORT_DIR / f"论文第二版_backup_before_se_enhance_{datetime.now():%Y%m%d_%H%M%S}.docx"
    shutil.copy2(DOC_PATH, backup)

    doc = Document(DOC_PATH)
    replace_text(doc)

    use_case = FIG_DIR / "figure_3_1_use_case.png"
    er = FIG_DIR / "figure_4_4_database_er.png"
    sequence = FIG_DIR / "figure_5_3_live_scan_sequence.png"

    anchor = find_following_table(doc, "表 3.1 系统功能需求")
    intro = paragraph(doc, "在功能需求基础上，系统用例可以进一步抽象为用户与外部服务之间的交互关系。用户侧主要完成账号登录、绘本管理、图片上传、实时识别、故事生成、语音朗读和历史查看；外部服务侧主要提供多模态识别与语音合成能力。")
    img = image_paragraph(doc, use_case)
    cap = caption(doc, "图3.1 系统用例图")
    test_intro = paragraph(doc, "为验证主要功能是否满足需求，本文整理了核心功能测试用例，如表 3.2 所示。测试用例覆盖用户登录、绘本上传、故事生成、实时识别、语音朗读和历史记录等主要业务路径。")
    test_cap = caption(doc, "表 3.2 功能测试用例表")
    test_tbl = table(
        doc,
        ["用例编号", "测试功能", "输入或操作", "预期结果", "测试结论"],
        [
            ["TC-01", "用户注册与登录", "输入用户名、邮箱和密码后登录", "返回访问令牌并进入系统页面", "通过"],
            ["TC-02", "绘本创建", "填写绘本标题并创建绘本", "绘本列表出现新增记录", "通过"],
            ["TC-03", "图片上传", "选择多张绘本页面图片上传", "图片按页序保存并可查询", "通过"],
            ["TC-04", "普通故事生成", "选择绘本并提交生成请求", "返回完整故事文本和任务状态", "通过"],
            ["TC-05", "实时识别", "启动摄像头并识别当前页面", "返回当前页讲述文本和识别状态", "通过"],
            ["TC-06", "连续页管理", "连续识别多页并切换总故事", "当前页与总故事文本分别保留", "通过"],
            ["TC-07", "语音朗读", "点击朗读当前讲述", "生成音频地址并可播放", "通过"],
            ["TC-08", "历史记录", "保存实时扫描故事后查看历史", "历史列表显示保存的故事记录", "通过"],
        ],
    )
    move_after(anchor, [intro._p, img._p, cap._p, test_intro._p, test_cap._p, test_tbl._tbl])

    anchor = find_paragraph(doc, "图4.1 系统总体架构")._p
    d1 = paragraph(doc, "从部署角度看，系统采用 Docker Compose 组织应用服务、MySQL 数据库和 Redis 缓存。应用容器负责运行 FastAPI 服务并对外暴露 8001 端口，MySQL 存储用户、绘本、图片和故事等结构化数据，Redis 用于缓存和实时扫描会话。上传图片、语音文件和日志通过宿主机目录挂载保存，便于云端部署和后续排查。")
    d2 = paragraph(doc, "这种部署结构的优点是组件边界清晰、迁移成本较低：本地开发可以使用 SQLite 或本地 MySQL，云端部署时通过环境变量切换数据库地址、Redis 地址、模型密钥和 TTS 配置。")
    move_after(anchor, [d1._p, d2._p])

    anchor = find_following_table(doc, "表 4.1 数据实体说明")
    db_intro = paragraph(doc, "除实体关系外，数据库表结构还需要说明字段、类型和约束，便于体现系统持久化设计。核心数据表结构如表 4.2 所示。")
    db_cap = caption(doc, "表 4.2 核心数据库表结构")
    db_tbl = table(
        doc,
        ["数据表", "关键字段", "字段类型", "约束或索引", "说明"],
        [
            ["users", "id, username, email, password_hash, created_at", "Integer / String / DateTime", "id 主键；username、email 唯一索引", "保存用户账号、邮箱、密码哈希和创建时间"],
            ["books", "id, user_id, title, cover_image, created_at", "Integer / String / DateTime", "user_id 外键并建立索引", "保存绘本元数据及所属用户"],
            ["book_images", "id, book_id, image_path, image_order, created_at", "Integer / String / DateTime", "book_id 外键并建立索引", "保存绘本页面图片路径和页序"],
            ["stories", "id, book_id, user_id, prompt, image_analysis, story_content, created_at", "Integer / Text / DateTime", "book_id、user_id 外键并建立索引", "保存故事文本、提示词和图像分析结果"],
        ],
    )
    er_intro = paragraph(doc, "各数据表之间的关系如图4.2 所示。用户与绘本、故事记录之间是一对多关系；绘本与绘本图片、故事记录之间也是一对多关系。")
    er_img = image_paragraph(doc, er, width=Inches(5.95))
    er_cap = caption(doc, "图4.2 数据库实体关系图")
    move_after(anchor, [db_intro._p, db_cap._p, db_tbl._tbl, er_intro._p, er_img._p, er_cap._p])

    anchor = find_following_table(doc, "表 5.1 核心接口说明")
    api_intro = paragraph(doc, "核心接口除路径和方法外，还需要明确主要请求字段和响应字段，便于说明后端实现不是简单页面调用。接口字段设计如表 5.2 所示。")
    api_cap = caption(doc, "表 5.2 核心接口请求与响应字段")
    api_tbl = table(
        doc,
        ["接口", "主要请求字段", "主要响应字段", "说明"],
        [
            ["/api/users/login", "username_or_email, password", "access_token, token_type, user", "完成用户认证并返回访问令牌"],
            ["/api/books", "title, cover_image", "book_id, title, created_at", "创建或查询当前用户绘本"],
            ["/api/images/{book_id}/images/upload", "files, image_order", "image_path, image_order, book_id", "上传并保存绘本页面图片"],
            ["/api/stories/generate", "book_id, prompt, style, target_age", "story_content, image_analysis, quality", "根据整本绘本生成故事"],
            ["/api/stories/scan", "image, session_id, mode, crop, style, target_age", "current_text, total_story, analysis, timing", "识别当前页并返回讲述结果"],
            ["/api/stories/scan/stream", "image, session_id, mode, crop, style, target_age", "delta, final, timing", "以 SSE 方式流式返回讲述文本"],
            ["/api/stories/tts", "text, voice, rate, volume", "audio_url, duration_seconds, provider", "将讲述文本转为语音文件"],
            ["/api/stories/scan/save", "session_id, title, story_content, images", "story_id, saved_at", "保存实时扫描生成的总故事"],
        ],
    )
    move_after(anchor, [api_intro._p, api_cap._p, api_tbl._tbl])

    anchor = find_paragraph(doc, "5.3 实时扫描会话与缓存实现")._p
    seq_intro = paragraph(doc, "实时识别链路涉及前端采集、后端扫描接口、实时讲述服务、多模态模型和 TTS 服务。为了说明各组件之间的调用顺序，本文将实时识别与朗读过程整理为图5.1。")
    seq_img = image_paragraph(doc, sequence, width=Inches(6.05))
    seq_cap = caption(doc, "图5.1 实时识别与朗读顺序图")
    move_after(anchor, [seq_intro._p, seq_img._p, seq_cap._p])

    # 论文表格统一按三线表处理；前两张通常是封面/元信息表，保留模板原样。
    for existing_table in doc.tables[2:]:
        make_three_line(existing_table)

    doc.save(DOC_PATH)
    print(f"updated={DOC_PATH}")
    print(f"backup={backup}")


if __name__ == "__main__":
    main()
